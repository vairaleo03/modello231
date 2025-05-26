from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import update
from app.database import get_db
from app.models.transcription_summaries import TranscriptionSummary
from pydantic import BaseModel
from app.routers.websocket_manager import websocket_manager;
from app.models.transcripts import Transcript
from app.models.transcription_summaries import TranscriptionSummary
from app.services.summarizer import generate_summary
from app.utils.post_processing import parse_summary_sections, compile_summary_docx
from datetime import datetime
from docx import Document

router = APIRouter()

class SummaryUpdateRequest(BaseModel):
    summary_text: str

# API che genera il riassunto della trascrizione
@router.post("/summary/start/{transcript_id}")
async def summarize_transcription(transcript_id: int, db: AsyncSession = Depends(get_db)):
    try: 
        result = await db.execute(select(Transcript).filter(Transcript.id == transcript_id))
        transcript = result.scalar_one_or_none()

        if not transcript:
            print(f"Trascrizione con ID {transcript_id} non trovata nel database.")
            raise HTTPException(status_code=404, detail="Trascrizione non trovata")

        if not transcript.transcript_text:
            raise HTTPException(status_code=400, detail="Testo della trascrizione mancante")

        summary = generate_summary(transcript.transcript_text)
                
        try:

            new_summary = TranscriptionSummary(
            transcript_id=transcript_id,
            summary_text=summary,
            created_at=datetime.utcnow()
            )

            db.add(new_summary)
            await db.commit()
            await db.refresh(new_summary)
        except Exception as e:
            print(f"❌ Errore durante il processo: {e}")

        return new_summary.id
    
    except Exception as e:
        print(f"❌ Eccezione nell'endpoint summary/start/ : {e}")
        raise HTTPException(status_code=500, detail=f"Errore durante il riassunto: {str(e)}")





# API che recupera un riassunto
@router.get("/summary/{summary_id}")
async def get_summary(summary_id: int, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(TranscriptionSummary).filter(TranscriptionSummary.id == summary_id))
    summary = result.scalar_one_or_none()

    if not summary:
        raise HTTPException(status_code=404, detail="Riassunto non trovato")

    return {
        "summary_id": summary.id,
        "transcript_id": summary.transcript_id,
        "summary_text": summary.summary_text,
        "created_at": summary.created_at
    }






# API che Salva automaticamente le modifiche al riassunto
@router.put("/summary/{summary_id}")
async def update_transcription(summary_id: int, request: SummaryUpdateRequest, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(TranscriptionSummary).filter(TranscriptionSummary.id == summary_id))
    summary = result.scalar_one_or_none()

    if not summary:
        raise HTTPException(status_code=404, detail="Riassunto non trovato")

    stmt = update(TranscriptionSummary).where(TranscriptionSummary.id == summary_id).values(summary_text=request.summary_text)
    await db.execute(stmt)
    await db.commit()
    await websocket_manager.send_notification("Modifiche salvate")
    return {"message": "Riassunto aggiornato con successo!"}





# Download di un riassunto come docx che include la formattazione del documento
# @router.post("/summary/{summary_id}/word")
# async def download_summary_docx(summary_id: int, db: AsyncSession = Depends(get_db)):
#     result = await db.execute(select(TranscriptionSummary).where(TranscriptionSummary.id == summary_id))
#     summary = result.scalar_one_or_none()

#     if not summary:
#         raise HTTPException(status_code=404, detail="Riassunto non trovato")

#     sections = parse_summary_sections(summary.summary_text)
#     template_path = "app/templates/template_riassunto.docx"
#     output_path = f"/tmp/riassunto_{summary_id}.docx"

#     compile_summary_docx(sections, template_path, output_path)

#     return FileResponse(
#         path=output_path,
#         filename=f"riassunto_{summary_id}.docx",
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )






# download di un riassunto come docx, in modo semplice, solo testo
@router.post("/summary/{summary_id}/word")
async def download_summary_word(summary_id: int, db: AsyncSession = Depends(get_db)):
    # Recupero riassunto dal DB
    result = await db.execute(select(TranscriptionSummary).filter_by(id=summary_id))
    summary = result.scalar_one_or_none()

    if not summary:
        raise HTTPException(status_code=404, detail="Riassunto non trovato")

    # Crea documento Word con il testo completo del riassunto
    doc = Document()
    doc.add_heading("Riassunto Trascrizione", level=1)
    doc.add_paragraph(summary.summary_text)

    # Salva in file temporaneo
    filename = f"riassunto_{summary_id}_{datetime.utcnow().timestamp()}.docx"
    filepath = f"/tmp/{filename}"
    doc.save(filepath)

    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )