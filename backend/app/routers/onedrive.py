from fastapi import APIRouter, Depends, HTTPException, Request, Response
from fastapi.responses import RedirectResponse, JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
import io
import os
import traceback  # ⭐ AGGIUNGI

from app.database import get_db
from app.services.onedrive_service import onedrive_service
from app.models.transcripts import Transcript
from app.models.transcription_summaries import TranscriptionSummary
from app.models.audio_files import AudioFile
from app.utils.post_processing import convert_html_to_word_template
from app.utils.session_manager import SessionManager

# Crea il router con prefisso
router = APIRouter(prefix="/onedrive", tags=["onedrive"])

# Frontend URL per reindirizzamenti
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")

@router.get("/auth")
async def authenticate_onedrive(response: Response):
    """
    Inizia il processo di autenticazione Microsoft, reindirizzando l'utente alla pagina di login.
    """
    try:
        print("🔐 Avvio autenticazione OneDrive...")
        auth_url = onedrive_service.get_auth_url()
        print(f"🔗 URL generato: {auth_url[:100]}...")
        return RedirectResponse(auth_url)
    except Exception as e:
        print(f"❌ Errore generazione auth URL: {str(e)}")
        print(f"❌ Traceback: {traceback.format_exc()}")
        return RedirectResponse(f"{FRONTEND_URL}?onedrive_auth=error&message=Errore_configurazione")

@router.get("/auth/callback")
async def onedrive_callback(request: Request, response: Response, code: str, state: str = None):
    """
    ⭐ CALLBACK AGGIORNATO SENZA COOKIE
    """
    try:
        print(f"📥 Callback ricevuto - Code: {code[:20]}..., State: {state}")
        
        # ⭐ VERIFICA STATE (importante per sicurezza)
        if state != "12345":
            print(f"❌ State mismatch: atteso '12345', ricevuto '{state}'")
            return RedirectResponse(f"{FRONTEND_URL}?onedrive_auth=error&message=State_mismatch")
        
        # ⭐ SCAMBIA CODICE CON TOKEN
        print("🔄 Scambio codice con token...")
        try:
            access_token, user_id = onedrive_service.get_token_from_code(code)
            print(f"✅ Token ottenuto per user: {user_id}")
            
        except Exception as token_error:
            print(f"❌ Errore scambio token: {str(token_error)}")
            error_msg = str(token_error).replace(" ", "_")
            return RedirectResponse(f"{FRONTEND_URL}?onedrive_auth=error&message={error_msg}")
        
        # ⭐ SALVA USANDO HEADER STRATEGY
        try:
            SessionManager.set_onedrive_user_id(response, request, user_id)
            print(f"✅ Headers preparati per user: {user_id}")
            
        except Exception as session_error:
            print(f"❌ Errore preparazione headers: {str(session_error)}")
        
        # ⭐ REDIRECT CON USER_ID NELL'URL
        success_url = f"{FRONTEND_URL}?onedrive_auth=success&user_id={user_id}&token={access_token[:20]}"
        print(f"🎯 Reindirizzamento con parametri: {success_url}")
        return RedirectResponse(success_url)
        
    except Exception as e:
        print(f"❌ Errore callback OneDrive: {str(e)}")
        error_msg = str(e).replace(" ", "_")
        return RedirectResponse(f"{FRONTEND_URL}?onedrive_auth=error&message={error_msg}")

@router.get("/auth/status")
async def check_auth_status(request: Request):
    """
    ⭐ STATUS CHECK MIGLIORATO
    """
    try:
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"🔍 Check auth status per user: {user_id}")
        
        if not user_id:
            print("❌ Nessun user_id in sessione")
            return {"authenticated": False, "reason": "no_session"}
        
        # Verifica se il token è ancora valido
        is_authenticated = onedrive_service.check_auth_status(user_id)
        print(f"✅ Status check result: {is_authenticated}")
        
        return {
            "authenticated": is_authenticated,
            "user_id": user_id if is_authenticated else None
        }
        
    except Exception as e:
        print(f"❌ Errore check auth status: {str(e)}")
        return {"authenticated": False, "error": str(e)}

@router.post("/auth/logout")
async def logout_onedrive(request: Request, response: Response):
    """
    Logout da OneDrive, rimuovendo i token e la sessione.
    """
    try:
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"👋 Logout OneDrive per user: {user_id}")
        
        # Rimuovi il token dalla cache, se presente
        if user_id and user_id in onedrive_service.token_cache:
            del onedrive_service.token_cache[user_id]
            print(f"✅ Token rimosso dalla cache per user: {user_id}")
        
        # Cancella la sessione
        SessionManager.clear_session(response)
        print("✅ Sessione cancellata")
        
        return {"success": True, "message": "Logout effettuato con successo"}
        
    except Exception as e:
        print(f"❌ Errore durante logout: {str(e)}")
        print(f"❌ Traceback: {traceback.format_exc()}")
        return {"success": False, "message": f"Errore durante logout: {str(e)}"}

# ⭐ ENDPOINT TEST CONNESSIONE UNIFICATO
@router.get("/test/connection")
async def test_onedrive_connection(request: Request):
    """
    Test endpoint per verificare la connessione OneDrive - versione unificata
    """
    try:
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"🧪 Test connessione per user: {user_id}")
        
        if not user_id:
            print("❌ Nessun user_id in sessione")
            return {"status": "error", "message": "Non autenticato - nessuna sessione"}
        
        # ⭐ Test token validity con logging
        print("🔍 Verifica validità token...")
        token = onedrive_service.get_valid_token(user_id)
        if not token:
            print("❌ Token non valido o scaduto")
            return {"status": "error", "message": "Token non valido o scaduto"}
            
        # ⭐ Test API call con timeout e logging dettagliato
        print("🌐 Test chiamata API Graph...")
        import requests
        headers = {"Authorization": f"Bearer {token}"}
        
        test_response = requests.get(
            "https://graph.microsoft.com/v1.0/me/drive/root",
            headers=headers,
            timeout=10
        )
        
        print(f"📊 API Response status: {test_response.status_code}")
        
        if test_response.status_code == 200:
            data = test_response.json()
            print("✅ Connessione OneDrive confermata")
            
            # ⭐ Informazioni unificate del drive
            drive_info = {
                "name": data.get("name"),
                "size": data.get("size"),
                "quota": data.get("quota", {}),
                "owner": None,
                "created_by": None
            }
            
            # ⭐ Gestione owner con fallback multipli
            if data.get("owner", {}).get("user", {}).get("displayName"):
                drive_info["owner"] = data.get("owner", {}).get("user", {}).get("displayName")
            elif data.get("createdBy", {}).get("user", {}).get("displayName"):
                drive_info["created_by"] = data.get("createdBy", {}).get("user", {}).get("displayName")
            
            print(f"📁 Drive info: {drive_info['name']}, Size: {drive_info.get('size', 'N/A')}")
            
            return {
                "status": "success", 
                "message": "Connessione OneDrive OK",
                "drive_info": drive_info
            }
        else:
            error_details = test_response.text[:200] if test_response.text else "Nessun dettaglio"
            print(f"❌ Errore API Graph: {test_response.status_code}")
            print(f"❌ Dettagli errore: {error_details}")
            
            return {
                "status": "error", 
                "message": f"Errore API Graph: {test_response.status_code}",
                "details": error_details
            }
            
    except Exception as e:
        print(f"❌ Eccezione durante test connessione: {str(e)}")
        print(f"❌ Traceback completo: {traceback.format_exc()}")
        return {"status": "error", "message": f"Eccezione test: {str(e)}"}

@router.post("/upload/transcription/{transcript_id}")
async def upload_transcription_to_onedrive(
    transcript_id: int, 
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    """
    ⭐ UPLOAD MIGLIORATO CON DEBUG
    """
    try:
        # Recupera l'ID utente dalla sessione
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"📤 Upload trascrizione {transcript_id} per user: {user_id}")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Autenticazione OneDrive richiesta")
        
        # Verifica che la trascrizione esista
        result = await db.execute(select(Transcript).filter(Transcript.id == transcript_id))
        transcription = result.scalar_one_or_none()
        
        if not transcription:
            raise HTTPException(status_code=404, detail="Trascrizione non trovata")
        
        print(f"✅ Trascrizione trovata: {len(transcription.transcript_text)} caratteri")
        
        # Genera il documento Word
        try:
            doc = convert_html_to_word_template(transcription.transcript_text)
            
            # Salva in memoria
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            file_content = file_stream.getvalue()
            
            print(f"✅ Documento Word generato: {len(file_content)} bytes")
            
        except Exception as doc_error:
            print(f"❌ Errore generazione documento: {str(doc_error)}")
            raise HTTPException(status_code=500, detail=f"Errore generazione documento: {str(doc_error)}")
        
        # Carica su OneDrive
        try:
            result = onedrive_service.upload_transcription(
                user_id, 
                transcript_id, 
                file_content
            )
            
            print(f"✅ Upload completato: {result.get('name')}")
            
            return {
                "success": True,
                "message": "Trascrizione caricata su OneDrive",
                "file_id": result.get("id"),
                "file_name": result.get("name"),
                "web_url": result.get("webUrl"),
                "size": len(file_content)
            }
            
        except Exception as upload_error:
            print(f"❌ Errore upload OneDrive: {str(upload_error)}")
            print(f"❌ Traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Errore upload: {str(upload_error)}")
        
    except HTTPException:
        # Rilancia le HTTPException
        raise
    except Exception as e:
        print(f"❌ Errore generale upload trascrizione: {str(e)}")
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Errore generale: {str(e)}")

@router.post("/upload/summary/{summary_id}")
async def upload_summary_to_onedrive(
    summary_id: int, 
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    """
    Carica un riassunto su OneDrive dell'utente.
    """
    try:
        # Recupera l'ID utente dalla sessione
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"📤 Upload riassunto {summary_id} per user: {user_id}")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Autenticazione OneDrive richiesta")
        
        # Verifica che il riassunto esista
        result = await db.execute(select(TranscriptionSummary).filter(TranscriptionSummary.id == summary_id))
        summary = result.scalar_one_or_none()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Riassunto non trovato")
            
        print(f"✅ Riassunto trovato: {len(summary.summary_text)} caratteri")
        
        # Crea documento Word dal riassunto
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Riassunto Trascrizione", level=1)
            
            # Formatta il testo in sezioni
            text_lines = summary.summary_text.split('\n')
            for line in text_lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.endswith(':'):
                    # È un titolo di sezione
                    doc.add_heading(line, level=2)
                elif line.startswith('- '):
                    # È un elemento di lista
                    doc.add_paragraph(line[2:], style='ListBullet')
                else:
                    # È un paragrafo normale
                    doc.add_paragraph(line)
            
            # Salva in memoria
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            file_content = file_stream.getvalue()
            
            print(f"✅ Documento Word riassunto generato: {len(file_content)} bytes")
            
        except Exception as doc_error:
            print(f"❌ Errore generazione documento riassunto: {str(doc_error)}")
            raise HTTPException(status_code=500, detail=f"Errore generazione documento: {str(doc_error)}")
        
        # Carica su OneDrive
        try:
            result = onedrive_service.upload_summary(
                user_id, 
                summary_id, 
                file_content
            )
            
            print(f"✅ Upload riassunto completato: {result.get('name')}")
            
            return {
                "success": True,
                "message": "Riassunto caricato su OneDrive",
                "file_id": result.get("id"),
                "file_name": result.get("name"),
                "web_url": result.get("webUrl"),
                "size": len(file_content)
            }
            
        except Exception as upload_error:
            print(f"❌ Errore upload riassunto OneDrive: {str(upload_error)}")
            print(f"❌ Traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Errore upload: {str(upload_error)}")
        
    except HTTPException:
        # Rilancia le HTTPException
        raise
    except Exception as e:
        print(f"❌ Errore generale upload riassunto: {str(e)}")
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Errore generale: {str(e)}")

@router.post("/upload/audio/{audio_id}")
async def upload_audio_to_onedrive(
    audio_id: int, 
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    """
    Carica un file audio su OneDrive dell'utente.
    """
    try:
        # Recupera l'ID utente dalla sessione
        user_id = SessionManager.get_onedrive_user_id(request)
        print(f"📤 Upload audio {audio_id} per user: {user_id}")
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Autenticazione OneDrive richiesta")
        
        # Verifica che il file audio esista
        result = await db.execute(select(AudioFile).filter(AudioFile.id == audio_id))
        audio = result.scalar_one_or_none()
        
        if not audio:
            raise HTTPException(status_code=404, detail="File audio non trovato")
            
        print(f"✅ File audio trovato: {audio.file_name}, {len(audio.file_data)} bytes")
        
        # Carica su OneDrive
        try:
            result = onedrive_service.upload_file(
                user_id, 
                audio.file_name, 
                audio.file_data, 
                "Modello231/Audio"
            )
            
            print(f"✅ Upload audio completato: {result.get('name')}")
            
            return {
                "success": True,
                "message": "File audio caricato su OneDrive",
                "file_id": result.get("id"),
                "file_name": result.get("name"),
                "web_url": result.get("webUrl"),
                "size": len(audio.file_data)
            }
            
        except Exception as upload_error:
            print(f"❌ Errore upload audio OneDrive: {str(upload_error)}")
            print(f"❌ Traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Errore upload: {str(upload_error)}")
        
    except HTTPException:
        # Rilancia le HTTPException
        raise
    except Exception as e:
        print(f"❌ Errore generale upload file audio: {str(e)}")
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Errore generale: {str(e)}")