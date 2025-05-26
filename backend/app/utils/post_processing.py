import re
from docx import Document
from bs4 import BeautifulSoup, NavigableString, Tag

######## converte i segmenti di una trascrizione in codice html per visualizzazione in editor tiptap nel frontend
def format_segments_html(segments: list[dict]) -> str:
    formatted = []
    speaker_map = {}
    speaker_index = 1
    current_speaker = None
    current_block = []

    segments_sorted = sorted(segments, key=lambda s: s["start"])

    for segment in segments_sorted:
        speaker = segment.get("speaker", "SPEAKER_??")
        text = segment.get("text", "").strip()
        if not text:
            continue

        # Mappa speaker
        if speaker not in speaker_map:
            speaker_map[speaker] = f"RELATORE {speaker_index}"
            speaker_index += 1

        speaker_label = speaker_map[speaker]

        # Inserisci ritorni a capo dopo ogni punto/virgola forte
        text = re.sub(r'(?<=[.!?])\s+', '<br>', text)

        # Se cambia relatore
        if current_speaker is not None and speaker != current_speaker:
            formatted.append(f"<p><strong>{speaker_map[current_speaker]}:</strong><br>{'<br>'.join(current_block)}</p>")
            current_block = []

        current_block.append(text)
        current_speaker = speaker

    if current_block and current_speaker:
        formatted.append(f"<p><strong>{speaker_map[current_speaker]}:</strong><br>{'<br>'.join(current_block)}</p>")

    return "\n".join(formatted)






####### converte il codice html di una trascrizione in testo word, per il downlaod in docx di una trascrizione

def convert_html_to_word(html_text):
    doc = Document()
    soup = BeautifulSoup(html_text, "html.parser")

    def process_node(node, paragraph, bold=False, italic=False, underline=False, highlight=None):
        if isinstance(node, NavigableString):
            run = paragraph.add_run(str(node))
            run.bold = bold if bold else None
            run.italic = italic if italic else None
            run.underline = underline if underline else None
            if highlight:
                try:
                    run.font.highlight_color = highlight
                except ValueError:
                    pass  # ignora colori non supportati da Word

        elif isinstance(node, Tag):
            new_bold = bold
            new_italic = italic
            new_underline = underline
            new_highlight = highlight

            if node.name == 'strong':
                new_bold = True
            if node.name == 'em':
                new_italic = True
            if node.name == 'u':
                new_underline = True

            if node.name == 'mark':
                new_highlight = "YELLOW"  # Word accetta: YELLOW, GREEN, PINK, BLUE, RED, GRAY, etc.

            if node.name == 'span':
                style = node.get("style", "")
                if "background-color" in style:
                    if "#ffff00" in style.lower():
                        new_highlight = "YELLOW"
                    elif "#ff0000" in style.lower():
                        new_highlight = "RED"
                    elif "#00ff00" in style.lower():
                        new_highlight = "GREEN"
                    elif "#00ffff" in style.lower():
                        new_highlight = "TURQUOISE"
                    # aggiungi altri colori HTML → Word se necessario

            if node.name == 'br':
                paragraph.add_run("\n")
            elif node.name in ['p', 'li']:
                paragraph = doc.add_paragraph()
                for child in node.children:
                    process_node(child, paragraph)
                doc.add_paragraph("")
            else:
                for child in node.children:
                    process_node(child, paragraph, new_bold, new_italic, new_underline, new_highlight)

    for element in soup.find_all(['p', 'li']):
        paragraph = doc.add_paragraph()
        for child in element.children:
            process_node(child, paragraph)
        doc.add_paragraph()

    return doc


def convert_html_to_word_template(text: str) -> Document:
    """
    Template base: prende testo semplice e lo inserisce nel documento Word.
    Ogni riga va a capo come paragrafo separato.
    (In progress)
    """
    doc = Document()
    
    # Evita errori se il testo è vuoto
    if not text:
        doc.add_paragraph("⚠️ Nessun contenuto disponibile.")
        return doc

    # Aggiunge ogni riga come paragrafo
    for line in text.split("\n"):
        doc.add_paragraph(line.strip())

    return doc


##### Estrae le sezioni da un riassunto
def parse_summary_sections(summary_text: str) -> dict:
    sections = {
        "temi_principali": "",
        "decisioni_prese": "",
        "responsabili_coinvolti": "",
        "prossimi_passi": ""
    }

    matches = re.split(r"\n?(\d\.\s.*)", summary_text)
    matches = [m.strip() for m in matches if m.strip()]

    for i in range(0, len(matches) - 1, 2):
        heading = matches[i].lower()
        content = matches[i + 1].strip()
        if "temi principali" in heading:
            sections["temi_principali"] = content
        elif "decisioni prese" in heading:
            sections["decisioni_prese"] = content
        elif "responsabili" in heading:
            sections["responsabili_coinvolti"] = content
        elif "prossimi passi" in heading:
            sections["prossimi_passi"] = content

    return sections






##### Compila il template word con le sezioni estratte dal riassunto
def compile_summary_docx(sections: dict, template_path: str, output_path: str):
    doc = Document(template_path)

    for p in doc.paragraphs:
        if "{{temi_principali}}" in p.text:
            p.text = p.text.replace("{{temi_principali}}", sections["temi_principali"])
        elif "{{decisioni_prese}}" in p.text:
            p.text = p.text.replace("{{decisioni_prese}}", sections["decisioni_prese"])
        elif "{{responsabili_coinvolti}}" in p.text:
            p.text = p.text.replace("{{responsabili_coinvolti}}", sections["responsabili_coinvolti"])
        elif "{{prossimi_passi}}" in p.text:
            p.text = p.text.replace("{{prossimi_passi}}", sections["prossimi_passi"])

    doc.save(output_path)
